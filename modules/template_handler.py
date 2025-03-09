import os
import re
import tempfile
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 同じディレクトリにあるllm_processingモジュールをインポート
from modules.llm_processing import get_azure_llm, get_groq_llm

def extract_placeholders(template_path):
    """
    Wordテンプレートからプレースホルダーを抽出する
    
    Args:
        template_path (str): テンプレートファイルのパス
        
    Returns:
        list: 抽出されたプレースホルダーのリスト
    """
    # テンプレートを読み込む
    doc = docx.Document(template_path)
    
    # プレースホルダーを保存するリスト
    placeholders = []
    
    # 正規表現パターン - {{任意の文字}}
    pattern = r'{{([^{}]+)}}'
    
    # すべてのパラグラフを検索
    for paragraph in doc.paragraphs:
        matches = re.findall(pattern, paragraph.text)
        placeholders.extend(matches)
    
    # すべてのテーブルを検索
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    matches = re.findall(pattern, paragraph.text)
                    placeholders.extend(matches)
    
    # 重複を除去
    unique_placeholders = list(set(placeholders))
    
    return unique_placeholders

def extract_info_with_llm(transcription, placeholders, api_choice='azure', model_type='llama3'):
    """
    文字起こしテキストからプレースホルダーに対応する情報をLLMで抽出する
    
    Args:
        transcription (str): 文字起こしテキスト
        placeholders (list): プレースホルダーのリスト
        api_choice (str): 使用するAPI ('azure' または 'groq')
        model_type (str): Groq使用時のモデルタイプ ('llama3' または 'gemma2')
        
    Returns:
        dict: プレースホルダーとその値のマッピング
    """
    # LLMを選択
    if api_choice.lower() == 'azure':
        llm = get_azure_llm()
    elif api_choice.lower() == 'groq':
        llm = get_groq_llm(model_type)
    else:
        raise ValueError(f"不明なAPI選択: {api_choice}")
    
    # プロンプトの作成
    prompt_text = f"""
文字起こしされた会議の内容から、以下の情報を抽出してください。
抽出できない場合は「情報なし」と記入してください。
情報は以下の形式で出力してください。

開催日時: [抽出された開催日時]
出席者: [抽出された出席者リスト]
...

以下は文字起こしの内容です:
{transcription}

抽出すべき情報:
"""
    
    for placeholder in placeholders:
        prompt_text += f"- {placeholder}\n"
    
    # LLMに質問
    from langchain_core.prompts import PromptTemplate
    from langchain_core.output_parsers import StrOutputParser
    
    prompt = PromptTemplate.from_template(prompt_text)
    chain = prompt | llm | StrOutputParser()
    
    response = chain.invoke({})
    
    # 抽出された情報をパースしてディクショナリに変換
    extracted_info = {}
    
    for placeholder in placeholders:
        # 正規表現パターン - "プレースホルダー: 任意の文字列"
        pattern = re.compile(f"{placeholder}:\\s*(.+?)(?=\\n\\w+:|$)", re.DOTALL)
        match = pattern.search(response)
        
        if match:
            # マッチした値を取得し、余分な空白を削除
            value = match.group(1).strip()
            extracted_info[placeholder] = value
        else:
            # 値が見つからない場合はデフォルト値
            extracted_info[placeholder] = "情報なし"
    
    return extracted_info

def process_template(template_path, transcription=None, summary=None, qa_data=None, api_choice='azure', model_type='llama3'):
    """
    Wordテンプレートを処理し、抽出した情報を挿入する
    
    Args:
        template_path (str): テンプレートファイルのパス
        transcription (str): 文字起こしテキスト
        summary (str): 要約テキスト
        qa_data (list): 質疑応答データのリスト
        api_choice (str): 使用するAPI ('azure' または 'groq')
        model_type (str): Groq使用時のモデルタイプ ('llama3' または 'gemma2')
        
    Returns:
        str: 生成された文書のパス
    """
    try:
        # テンプレートからプレースホルダーを抽出
        placeholders = extract_placeholders(template_path)
        print(f"抽出されたプレースホルダー: {placeholders}")
        
        # 必要な情報を準備
        replacements = {}
        
        # 「要約」プレースホルダーが存在し、summary引数が与えられている場合
        if '要約' in placeholders and summary:
            replacements['要約'] = summary
            # 処理済みのプレースホルダーを削除
            if '要約' in placeholders:
                placeholders.remove('要約')
        
        # 残りのプレースホルダーに対応する情報を文字起こしから抽出
        if transcription and placeholders:
            extracted_info = extract_info_with_llm(
                transcription, 
                placeholders, 
                api_choice=api_choice, 
                model_type=model_type
            )
            
            # 抽出された情報をreplacementsに追加
            replacements.update(extracted_info)
        
        # テンプレートを読み込む
        doc = docx.Document(template_path)
        
        # テンプレート内のプレースホルダーを置換
        replace_placeholders_in_document(doc, replacements)
        
        # 文字起こし全文を追加（特定のプレースホルダーがあれば置換、なければ末尾に追加）
        if transcription:
            if '文字起こし全文' in placeholders and '文字起こし全文' in replacements:
                # すでに処理済み
                pass
            else:
                # ドキュメントの末尾に追加
                doc.add_heading('文字起こし全文', level=1)
                doc.add_paragraph(transcription)
        
        # Q&Aデータがある場合はテーブルに追加
        if qa_data:
            add_qa_table(doc, qa_data)
        
        # 一時ファイルとして保存
        output_dir = 'static/uploads'
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        output_path = os.path.join(output_dir, 'meeting_minutes.docx')
        doc.save(output_path)
        
        return output_path
    
    except Exception as e:
        raise Exception(f"テンプレート処理中にエラーが発生しました: {str(e)}")

def replace_placeholders_in_document(doc, replacements):
    """
    ドキュメント内のプレースホルダーを置換する
    
    Args:
        doc (Document): docxドキュメント
        replacements (dict): プレースホルダーと置換値のマッピング
    """
    # すべてのパラグラフを処理
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    
    # すべてのテーブルを処理
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

def replace_text_in_paragraph(paragraph, replacements):
    """
    パラグラフ内のプレースホルダーを置換する
    
    Args:
        paragraph (Paragraph): docxパラグラフ
        replacements (dict): プレースホルダーと置換値のマッピング
    """
    for placeholder, value in replacements.items():
        if f"{{{{{placeholder}}}}}" in paragraph.text:
            # テキスト置換のための準備
            placeholder_text = f"{{{{{placeholder}}}}}"
            text = paragraph.text
            
            # テキストを置換
            new_text = text.replace(placeholder_text, value)
            
            # 古いランの保存
            old_runs = [run for run in paragraph.runs]
            
            # パラグラフをクリア
            for i in range(len(paragraph.runs)):
                p = paragraph._p
                p.remove(paragraph.runs[0]._r)
            
            # 新しいテキストを追加（スタイルを維持）
            if old_runs:
                # 最初のランのスタイルを使用
                run = paragraph.add_run(new_text)
                run.bold = old_runs[0].bold
                run.italic = old_runs[0].italic
                run.underline = old_runs[0].underline
                run.font.name = old_runs[0].font.name
                if old_runs[0].font.size:
                    run.font.size = old_runs[0].font.size
            else:
                paragraph.add_run(new_text)

def add_qa_table(doc, qa_data):
    """
    Q&Aテーブルをドキュメントに追加
    
    Args:
        doc (Document): docxドキュメント
        qa_data (list): 質疑応答データのリスト
    """
    # Q&Aセクションヘッダーを追加
    doc.add_heading("質疑応答", level=2)
    
    # Q&Aテーブルを作成
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # ヘッダー行を設定
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "質問"
    hdr_cells[1].text = "回答"
    
    # ヘッダー行のスタイルを設定
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
    
    # Q&Aデータを追加
    for qa in qa_data:
        row_cells = table.add_row().cells
        row_cells[0].text = qa.get('question', '')
        row_cells[1].text = qa.get('answer', '')
