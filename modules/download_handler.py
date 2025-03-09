import os
import json
import docx
from datetime import datetime

def generate_download(transcription, summary=None, qa_data=None, output_format='docx'):
    """
    ダウンロード用のファイルを生成する
    
    Args:
        transcription (str): 文字起こしテキスト
        summary (str): 要約テキスト
        qa_data (list): 質疑応答データのリスト
        output_format (str): 出力フォーマット ('docx', 'txt', 'json')
        
    Returns:
        str: 生成されたファイルのパス
    """
    # 出力ディレクトリを確保
    output_dir = 'static/uploads'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # タイムスタンプを生成
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # ファイル名と相対パスを生成
    filename = f'report_{timestamp}.{output_format}'
    relative_path = os.path.join('static/uploads', filename)
    
    # 絶対パスを生成 (ファイル保存用)
    try:
        # 現在のファイルの場所を基準にする
        base_dir = os.path.dirname(os.path.abspath(__file__))
        parent_dir = os.path.dirname(base_dir)
        abs_path = os.path.join(parent_dir, 'static/uploads', filename)
    except:
        # 絶対パスの生成に失敗した場合は相対パスをそのまま使用
        abs_path = relative_path
    
    print(f"ファイル生成情報:")
    print(f"  出力フォーマット: {output_format}")
    print(f"  タイムスタンプ: {timestamp}")
    print(f"  ファイル名: {filename}")
    print(f"  相対パス: {relative_path}")
    print(f"  絶対パス: {abs_path}")
    
    if output_format.lower() == 'docx':
        return generate_docx(transcription, summary, qa_data, abs_path)
    elif output_format.lower() == 'txt':
        return generate_txt(transcription, summary, qa_data, abs_path)
    elif output_format.lower() == 'json':
        return generate_json(transcription, summary, qa_data, abs_path)
    else:
        raise ValueError(f"サポートされていない出力フォーマット: {output_format}")

def generate_docx(transcription, summary, qa_data, output_path):
    """
    Wordドキュメントを生成
    
    Args:
        transcription (str): 文字起こしテキスト
        summary (str): 要約テキスト
        qa_data (list): 質疑応答データのリスト
        output_path (str): 出力ファイルパス
        
    Returns:
        str: 生成されたファイルの相対パス
    """
    # 新しいドキュメントを作成
    doc = docx.Document()
    
    # タイトルを追加
    doc.add_heading('音声文字起こし報告書', 0)
    
    # 生成日時を追加
    doc.add_paragraph(f'生成日時: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
    
    # 以下略...
    
    # ファイルを保存
    print(f"Word文書の保存先: {output_path}")
    doc.save(output_path)
    
    # 相対パスを返す (static/uploads/filename.docx 形式)
    relative_path = 'static/uploads/' + os.path.basename(output_path)
    return relative_path
    
def generate_txt(transcription, summary, qa_data, output_dir, timestamp):
    """
    テキストファイルを生成
    
    Args:
        transcription (str): 文字起こしテキスト
        summary (str): 要約テキスト
        qa_data (list): 質疑応答データのリスト
        output_dir (str): 出力ディレクトリ
        timestamp (str): タイムスタンプ
        
    Returns:
        str: 生成されたファイルのパス
    """
    content = []
    
    # タイトルを追加
    content.append("音声文字起こし報告書")
    content.append("="*40)
    
    # 生成日時を追加
    content.append(f'生成日時: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
    content.append("")
    
    # 要約セクションを追加（存在する場合）
    if summary:
        content.append("【要約】")
        content.append("-"*40)
        content.append(summary)
        content.append("")
    
    # 文字起こしセクションを追加
    content.append("【文字起こし全文】")
    content.append("-"*40)
    content.append(transcription)
    content.append("")
    
    # Q&Aセクションを追加（存在する場合）
    if qa_data and len(qa_data) > 0:
        content.append("【質疑応答】")
        content.append("-"*40)
        
        for i, qa in enumerate(qa_data, 1):
            content.append(f"Q{i}: {qa.get('question', '')}")
            content.append(f"A{i}: {qa.get('answer', '')}")
            content.append("")
    
    # ファイルを保存
    output_path = os.path.join(output_dir, f'report_{timestamp}.txt')
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(content))
    
    return os.path.join(os.path.basename(output_dir), f'report_{timestamp}.txt')

def generate_json(transcription, summary, qa_data, output_dir, timestamp):
    """
    JSONファイルを生成
    
    Args:
        transcription (str): 文字起こしテキスト
        summary (str): 要約テキスト
        qa_data (list): 質疑応答データのリスト
        output_dir (str): 出力ディレクトリ
        timestamp (str): タイムスタンプ
        
    Returns:
        str: 生成されたファイルのパス
    """
    # JSONデータを作成
    data = {
        'timestamp': datetime.now().isoformat(),
        'transcription': transcription,
        'summary': summary if summary else "",
        'qa_data': qa_data if qa_data else []
    }
    
    # ファイルを保存
    output_path = os.path.join(output_dir, f'report_{timestamp}.json')
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    
    return os.path.join(os.path.basename(output_dir), f'report_{timestamp}.json')
